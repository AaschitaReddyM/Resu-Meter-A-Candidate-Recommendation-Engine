# ---Source code for Streamlit application---

#essential because streamlit needs a .py file to run
%%writefile app.py

#importing all the libraries previously installed, making their functions available to the script
import streamlit as st
import numpy as np
from sentence_transformers import SentenceTransformer
import google.generativeai as genai
import os
import time
import PyPDF2
import docx
from io import BytesIO
import spacy
import pandas as pd
import plotly.express as px


# ---Loading AI models and API Configuration---
genai.configure(api_key=os.environ.get('GOOGLE_API_KEY'))

@st.cache_resource #using a decorator to indicate streamlit to be cautious about the next declared function.
def load_sentence_model():
#loading the sentence-transformer model and cacheing it for faster execution
    return SentenceTransformer('all-MiniLM-L6-v2')

@st.cache_resource
def load_spacy_model():
#loading the spacy model and cacheing it
    try:
        return spacy.load("en_core_web_sm")
    except OSError:
        print("Downloading the spaCy model")
        spacy.cli.download("en_core_web_sm")
        return spacy.load("en_core_web_sm")

#loading the models
model= load_sentence_model()
nlp= load_spacy_model()

# ---Helper Functions---
#these are the specialised funcions designed to perform the core logic of the application

#function to read uploaded files
def extract_text_from_file(upload_file):
    try:
        if upload_file.type =="text/plain":
            return uploaded_file.read().decode("utf-8")
            # .read(): Method that reads a file's content as raw bytes.
            # .decode(): Method that converts bytes into a readable string.
        elif uploaded_file.type== "application/pdf":
            pdf_file= BytesIO(uploaded_file.getvalue()) # .getvalue(): Method that retrieves the raw byte content of the uploaded file.
            # BytesIO(): Class that creates an in-memory file-like object from bytes.
            pdf_reader= PyPDF2.PdfReader(pdf_file) # PyPDF2.PdfReader(): Class that creates an object to read and parse a PDF file.
            text = ""
            for page in pdf_reader.pages:  # .pages: Attribute that contains a list of all page objects within the PDF.
                text += page.extract_text() or ""
            return text
        elif uploaded_file.type== "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc= docx.Document(BytesIO(uploaded_file.getvalue())) # docx.Document(): Class that creates an object to read and parse a .docx file.
            text= ""
            for para in doc.paragraphs: # .paragraphs: Attribute containing a list of all paragraph objects in the document.
                  text += para.text + "\\n"
            return text
    except Exception as e:
        st.error(f"Error reading file {uploaded_file.name}: {e}")
    return None

# Calculates the cosine similarity score between two numerical vectors.
def cosine_similarity(a, b):
    dot_product= np.dot(a, b) # np.dot(): Function that calculates the dot product of two vectors.
    norm_a= np.linalg.norm(a) # np.linalg.norm(): Function that calculates the length (norm) of a vector.
    norm_b= np.linalg.norm(b)
    return 0.0 if norm_a ==0 or norm_b==0 else dot_product / (norm_a * norm_b)

# Generates an AI summary by sending a prompt to the Google Gemini API.
def get_gemini_summary(job_description, resume_text):
    prompt = (
        f"Based on the following job description and candidate resume, "
        f"write a concise summary (1-2 sentences) explaining why this candidate "
        f"is a great fit for the role.\\n\\n"
        f"Job Description:\\n{job_description}\\n\\n"
        f"Candidate Resume:\\n{resume_text}"
    )
    try:
        model_gemini = genai.GenerativeModel('gemini-2.5-flash') #using the latest 2.5 flash model for higher accuracy and reliability
        response = model_gemini.generate_content(prompt)
        return response.text or "Could not generate a summary from the API response."
    except Exception as e:
        return f"Error generating summary: {str(e)}"

def extract_keywords(text):
    """Extracts nouns and proper nouns as keywords, filtering for actual words."""
    doc = nlp(text.lower())
    keywords = {
        token.text for token in doc
        if token.pos_ in ['NOUN', 'PROPN'] and token.is_alpha and len(token.text) > 1
    }
    return keywords

# --- Streamlit UI ---
st.set_page_config(page_title="Resu-Meter", page_icon="🎯", layout="wide")

st.markdown("""
    <style>
    .big-font { font-size:40px !important; font-weight: bold; color: #4A90E2; }
    .stButton>button { width: 100%; font-size: 1.25rem; font-weight: bold; }
    </style>
""", unsafe_allow_html=True)

st.markdown('<p class="big-font">Resu-Meter: A Candidate Recommendation Engine</p>', unsafe_allow_html=True)
st.write("Hi Recruiter! Find the best candidates by uploading resumes or pasting text. The engine will rank them and highlight key skills.")

with st.form("recommendation_form"):
    st.subheader("1. Job Description")
    job_description = st.text_area("Please paste the job description here...", height=200, key="job_desc")

    st.subheader("2. Candidate Resumes")
    tab1, tab2 = st.tabs(["📄 Upload Resumes", "✍️ Paste Resumes"])

    with tab1:
        uploaded_files = st.file_uploader(
            "Choose resume files (accepted formats: .pdf, .docx, .txt)",
            accept_multiple_files=True,
            type=['pdf', 'docx', 'txt']
        )
    with tab2:
        if 'candidates' not in st.session_state:
            st.session_state.candidates = [{'name': '', 'resume': ''}]
        for i, candidate in enumerate(st.session_state.candidates):
            st.text_input(f"Candidate Name or ID", key=f"name_{i}", placeholder=f"Candidate {i+1} Name")
            st.text_area(f"Resume Text", height=150, key=f"resume_{i}", placeholder="Please paste resume text here...")

    submit_button = st.form_submit_button("🧑‍💼👍 Generate Recommendations")

# --- Recommendation Logic ---
if submit_button:
    all_candidates = []
    if uploaded_files:
        for uploaded_file in uploaded_files:
            resume_text = extract_text_from_file(uploaded_file)
            if resume_text:
                all_candidates.append({'name': uploaded_file.name, 'resume': resume_text})

    pasted_candidates = [{'name': st.session_state[f'name_{i}'], 'resume': st.session_state[f'resume_{i}']} for i in range(len(st.session_state.candidates))]
    valid_pasted_candidates = [c for c in pasted_candidates if c['name'] and c['resume']]
    all_candidates.extend(valid_pasted_candidates)

    if not job_description or not all_candidates:
        st.error("❗ Please provide a job description and at least one resume.")
    else:
        st.subheader("🏆 Top Candidate Recommendations")
        with st.spinner('Analyzing resumes, generating recommendations, and building visualizations...'):
            job_embedding = model.encode(job_description)
            job_keywords = extract_keywords(job_description)
            recommendations = []

            for candidate in all_candidates:
                resume_embedding = model.encode(candidate['resume'])
                similarity = cosine_similarity(job_embedding, resume_embedding)
                recommendations.append({
                    'name': candidate['name'],
                    'similarity': float(similarity),
                    'resume': candidate['resume']
                })

            recommendations.sort(key=lambda x: x['similarity'], reverse=True)
            top_recommendations = recommendations[:10]

            # --- Visualization Part ---
            df = pd.DataFrame(top_recommendations)
            if not df.empty:
                fig = px.bar(df, x='similarity', y='name', orientation='h', title='Top Candidate Scores',
                             text='similarity', labels={'similarity': 'Relevance Score', 'name': 'Candidate'})
                fig.update_layout(yaxis={'categoryorder':'total ascending'}, uniformtext_minsize=8, uniformtext_mode='hide')
                fig.update_traces(texttemplate='%{text:.2%}', textposition='outside')
                st.plotly_chart(fig, use_container_width=True)

        st.success("🎉 Recommendations generated!")
        for i, rec in enumerate(top_recommendations):
            st.markdown(f"### **{i+1}. {rec['name']}**")
            st.progress(rec['similarity'], text=f"**Relevance Score: {rec['similarity']:.2%}**")

            # --- Keyword Matching Part ---
            resume_keywords = extract_keywords(rec['resume'])
            matched_keywords = job_keywords.intersection(resume_keywords)
            if matched_keywords:
                st.markdown(f"**🔑 Matched Keywords:** `{'`, `'.join(sorted(list(matched_keywords)))}`") #avoiding the identification of special characters instead of keywords

            with st.expander(" Show AI-Generated Summary"):
                with st.spinner('Generating summary...'):
                    summary = get_gemini_summary(job_description, rec['resume'])
                    st.write(summary)